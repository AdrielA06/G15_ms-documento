from typing import Dict, Any, Optional
from jinja2 import Environment, FileSystemLoader
import pdfkit
from docx import Document
from docx.shared import Inches, Pt
from odf.opendocument import OpenDocumentText
from odf.text import P
from datetime import datetime
import os

class DocumentoService:    
    def __init__(self, template_dir: str = 'app/template'):
        self.template_dir = template_dir
        self.jinja_env = Environment(loader=FileSystemLoader(template_dir))
        
    def generar_pdf(self, template_name: str, data: Dict[str, Any], 
                    output_path: str, config: Optional[Dict] = None) -> str:
        html_content = self._render_template(template_name, data)

#REESTRUCTURAR ESTA PARTE 
        pdf_config = config or {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'enable-local-file-access': None
        }
        
        pdfkit.from_string(html_content, output_path, options=pdf_config)
        
        return output_path
    
    def generar_docx(self, template_name: str, data: Dict[str, Any], 
                     output_path: str) -> str:
        html_content = self._render_template(template_name, data)
        document = Document()
        titulo = document.add_heading(data.get('titulo', 'Documento'), 0)
        self._html_to_docx(html_content, document, data)        
        document.save(output_path)
        return output_path
    
    def generar_odt(self, template_name: str, data: Dict[str, Any], 
                    output_path: str) -> str:
        html_content = self._render_template(template_name, data)
        textdoc = OpenDocumentText()
        self._html_to_odt(html_content, textdoc, data)
        textdoc.save(output_path)
        
        return output_path
    
    def _render_template(self, template_name: str, data: Dict[str, Any]) -> str:
        template = self.jinja_env.get_template(template_name)
        return template.render(**data)
    
    def _html_to_docx(self, html_content: str, document: Document, data: Dict[str, Any]):
        p = document.add_paragraph()
        p.add_run(f"Legajo: {data.get('legajo', 'N/A')}").bold = True
        
        document.add_paragraph(f"Nombre: {data.get('nombre_completo', 'N/A')}")
        document.add_paragraph(f"DNI: {data.get('numero_documento', 'N/A')}")
        document.add_paragraph(f"Especialidad: {data.get('especialidad', 'N/A')}")
        document.add_paragraph(f"Año de Cursado: {data.get('anio_cursado', 'N/A')}")
    
    def _html_to_odt(self, html_content: str, textdoc: OpenDocumentText, data: Dict[str, Any]):
        textdoc.text.addElement(P(text=f"Legajo: {data.get('legajo', 'N/A')}"))
        textdoc.text.addElement(P(text=f"Nombre: {data.get('nombre_completo', 'N/A')}"))
        textdoc.text.addElement(P(text=f"DNI: {data.get('numero_documento', 'N/A')}"))
        textdoc.text.addElement(P(text=f"Especialidad: {data.get('especialidad', 'N/A')}"))
        textdoc.text.addElement(P(text=f"Año de Cursado: {data.get('anio_cursado', 'N/A')}"))